"""
Extract structured text data from bid response Word (.docx) files.

Scans all tables in a Word document, classifies them by content patterns,
and outputs structured JSON files for company info, personnel, qualifications,
performance, and cost breakdown.

Usage:
    python extract_text.py <docx_path> [--output-dir data] [--index index.json]

The --index option enables cross-referencing with an existing image index.
"""
import argparse
import json
import os
import re
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn


PLACEHOLDER_RE = re.compile(r'【.*?】')


def is_placeholder(value: str) -> bool:
    """Check if a value is or contains an unfilled placeholder."""
    if not value:
        return False
    return bool(PLACEHOLDER_RE.search(value.strip()))


def clean_value(value: str):
    """Return value or None if it's a placeholder."""
    if not value or not value.strip():
        return None
    v = value.strip()
    if is_placeholder(v):
        return None
    # Also treat common placeholder patterns
    if v in ('——', '—', '-', '/'):
        return None
    return v


def table_to_rows(table) -> list[list[str]]:
    """Extract all rows from a table, handling merged cells."""
    rows = []
    for row in table.rows:
        cells = []
        seen_texts = set()
        for cell in row.cells:
            text = cell.text.strip()
            # Deduplicate merged cells (they appear multiple times)
            cell_id = id(cell._tc)
            if cell_id in seen_texts:
                continue
            seen_texts.add(cell_id)
            cells.append(text)
        rows.append(cells)
    return rows


def get_table_contexts(doc) -> dict[int, str]:
    """
    Build a mapping of table index -> preceding paragraph context.
    Walks the document body XML to find paragraphs before each table.
    """
    contexts = {}
    table_index = 0
    last_paragraphs = []

    for child in doc.element.body:
        if child.tag == qn('w:p'):
            # Extract text from paragraph
            texts = []
            for r in child.iter(qn('w:t')):
                if r.text:
                    texts.append(r.text)
            text = ''.join(texts).strip()
            if text:
                last_paragraphs.append(text)
                # Keep last 5 paragraphs for context
                if len(last_paragraphs) > 5:
                    last_paragraphs = last_paragraphs[-5:]
        elif child.tag == qn('w:tbl'):
            contexts[table_index] = '\n'.join(last_paragraphs)
            table_index += 1

    return contexts


def classify_table(rows: list[list[str]], context: str) -> str:
    """
    Classify a table by its content patterns.
    Returns one of: company, legal_rep, authorized_rep, personnel_detail,
    team_summary, performance, qualification, bank, enterprise_scale,
    cost_breakdown, pricing, unknown
    """
    if not rows or len(rows) < 2:
        return 'unknown'

    header = rows[0]
    first_col_vals = [r[0] for r in rows if r]
    all_text = ' '.join(' '.join(r) for r in rows)
    ctx_lower = context.lower() if context else ''

    # 2-column key-value table
    if len(header) == 2:
        keys = set(first_col_vals)

        # Company overview table: has 供应商名称 + 统一社会信用代码
        if '供应商名称' in keys and '统一社会信用代码' in keys:
            if '开户银行' in keys and '银行账号' in keys:
                # Could be a combined company+bank table
                if '注册地址' in keys:
                    return 'company'  # Full company table with bank fields
                else:
                    return 'bank'
            if '注册地址' in keys or '注册资金' in keys:
                return 'company'

        # Bank table
        if '开户银行' in keys and '银行账号' in keys:
            return 'bank'

        # Enterprise scale table
        if '企业类型' in keys or ('从业人员' in keys and '企业名称' in keys):
            return 'enterprise_scale'

        # Personnel detail (name + education/cert/years)
        if '姓名' in keys:
            has_edu = '学历' in keys
            has_cert = any('职称' in k or '认证' in k for k in keys)
            has_years = any('年限' in k for k in keys)
            has_id = any('身份证' in k for k in keys)

            # Determine person type from context
            # "授权代表" / "授权委托" patterns take priority over "法定代表人"
            # because authorization letters mention both: "法定代表人XX授权委托YY"
            is_auth_context = ('授权代表' in context or '授权委托' in context
                               or '被授权人' in context)
            # Only classify as legal_rep if context mentions 法定代表人 WITHOUT
            # also being an authorization context (e.g., "法定代表人身份证明")
            is_legal_context = ('法定代表人' in context and not is_auth_context)

            if has_edu or has_cert or has_years:
                if is_auth_context:
                    return 'authorized_rep'
                if is_legal_context:
                    return 'legal_rep'
                return 'personnel_detail'

            if has_id and '性别' in keys:
                if is_auth_context:
                    return 'authorized_rep'
                if is_legal_context:
                    return 'legal_rep'
                return 'personnel_detail'

        # Simple team stats (人数 + 职责)
        if '人数' in keys and ('角色职责' in keys or '职责' in keys):
            return 'team_summary_single'

    # Multi-column tables
    header_text = ' '.join(header)

    # Performance table
    if ('项目名称' in header_text and ('采购人' in header_text or '甲方' in header_text)):
        return 'performance'
    if '合同金额' in header_text and '项目名称' in header_text:
        return 'performance'

    # Team summary table
    if '岗位' in header_text and '人数' in header_text:
        return 'team_summary'

    # Cost breakdown
    if '费用项目' in header_text and ('金额' in header_text or '费用' in header_text):
        return 'cost_breakdown'

    # Pricing table
    if '报价' in header_text or ('单价' in header_text and '总价' in header_text):
        return 'pricing'

    # Technical response table (skip - project-specific)
    if '需求内容' in header_text and '是否响应' in header_text:
        return 'tech_response'

    # Qualification table
    if '证书名称' in header_text or '认证类型' in header_text:
        return 'qualification'

    return 'unknown'


def extract_kv_table(rows: list[list[str]]) -> dict[str, str]:
    """Extract key-value pairs from a 2-column table."""
    result = {}
    for row in rows:
        if len(row) >= 2 and row[0]:
            result[row[0].strip()] = row[1].strip() if len(row) > 1 else ''
    return result


def extract_company(tables_data: list[tuple]) -> dict:
    """Extract company information from classified tables."""
    company = {
        'name': None, 'credit_code': None, 'registered_address': None,
        'office_address': None, 'registered_capital': None, 'company_type': None,
        'established_date': None, 'postal_code': None, 'email': None,
        'phone': None, 'fax': None,
        'legal_representative': {
            'name': None, 'gender': None, 'title': None, 'id_number': None
        },
        'authorized_representative': {
            'name': None, 'gender': None, 'age': None, 'title': None,
            'id_number': None, 'phone': None, 'email': None
        },
        'bank': {
            'bank_name': None, 'account_number': None, 'bank_code': None
        },
        'enterprise_scale': {
            'type': None, 'employees': None, 'revenue': None,
            'total_assets': None, 'industry': None
        }
    }

    placeholders = []

    for ttype, rows, context, ti in tables_data:
        kv = extract_kv_table(rows)

        if ttype == 'company':
            field_map = {
                '供应商名称': 'name', '企业名称': 'name',
                '统一社会信用代码': 'credit_code',
                '注册地址': 'registered_address',
                '办公地址': 'office_address',
                '注册资金': 'registered_capital', '注册资本': 'registered_capital',
                '企业性质': 'company_type', '企业类型': 'company_type',
                '成立日期': 'established_date',
                '邮政编码': 'postal_code', '邮编': 'postal_code',
                '电子邮箱': 'email', '邮箱': 'email',
                '联系电话': 'phone', '电话': 'phone',
                '传真': 'fax',
                '法定代表人': ('legal_representative', 'name'),
                '联系人': '_contact_person',
            }
            for k, v in kv.items():
                for fk, target in field_map.items():
                    if fk in k:
                        if is_placeholder(v):
                            placeholders.append({
                                'field': target if isinstance(target, str) else f'{target[0]}.{target[1]}',
                                'original': v, 'table_index': ti
                            })
                            break
                        val = clean_value(v)
                        if isinstance(target, tuple):
                            company[target[0]][target[1]] = val
                        elif target.startswith('_'):
                            pass  # skip internal fields
                        else:
                            if company.get(target) is None:  # don't overwrite
                                company[target] = val
                        break

            # Bank fields might be in company table
            if '开户银行' in kv:
                bv = clean_value(kv['开户银行'])
                if bv:
                    company['bank']['bank_name'] = bv
                elif is_placeholder(kv['开户银行']):
                    placeholders.append({'field': 'bank.bank_name', 'original': kv['开户银行'], 'table_index': ti})
            if '银行账号' in kv:
                bv = clean_value(kv['银行账号'])
                if bv:
                    company['bank']['account_number'] = bv
                elif is_placeholder(kv['银行账号']):
                    placeholders.append({'field': 'bank.account_number', 'original': kv['银行账号'], 'table_index': ti})

        elif ttype == 'legal_rep':
            rep_map = {
                '姓名': 'name', '性别': 'gender', '职务': 'title',
                '身份证号码': 'id_number', '身份证号': 'id_number',
            }
            for k, v in kv.items():
                for fk, target in rep_map.items():
                    if fk in k:
                        if is_placeholder(v):
                            placeholders.append({'field': f'legal_representative.{target}', 'original': v, 'table_index': ti})
                            break
                        company['legal_representative'][target] = clean_value(v)
                        break

        elif ttype == 'authorized_rep':
            rep_map = {
                '姓名': 'name', '性别': 'gender', '年龄': 'age',
                '职务': 'title', '身份证号码': 'id_number', '身份证号': 'id_number',
                '联系电话': 'phone', '电话': 'phone',
                '电子邮箱': 'email', '邮箱': 'email',
            }
            for k, v in kv.items():
                for fk, target in rep_map.items():
                    if fk in k:
                        if is_placeholder(v):
                            placeholders.append({'field': f'authorized_representative.{target}', 'original': v, 'table_index': ti})
                            break
                        company['authorized_representative'][target] = clean_value(v)
                        break

        elif ttype == 'bank':
            bank_map = {
                '开户银行': 'bank_name', '银行账号': 'account_number',
                '开户行联行号': 'bank_code', '联行号': 'bank_code',
            }
            for k, v in kv.items():
                for fk, target in bank_map.items():
                    if fk in k:
                        if is_placeholder(v):
                            placeholders.append({'field': f'bank.{target}', 'original': v, 'table_index': ti})
                            break
                        company['bank'][target] = clean_value(v)
                        break

        elif ttype == 'enterprise_scale':
            scale_map = {
                '企业类型': 'type', '企业规模': 'type',
                '从业人员': 'employees', '员工人数': 'employees',
                '营业收入': 'revenue',
                '资产总额': 'total_assets',
                '所属行业': 'industry', '行业': 'industry',
            }
            for k, v in kv.items():
                for fk, target in scale_map.items():
                    if fk in k:
                        if is_placeholder(v):
                            placeholders.append({'field': f'enterprise_scale.{target}', 'original': v, 'table_index': ti})
                            break
                        company['enterprise_scale'][target] = clean_value(v)
                        break

    return company, placeholders


def extract_personnel(tables_data: list[tuple]) -> dict:
    """Extract personnel information from classified tables."""
    personnel = []
    team_summary = []
    placeholders = []

    for ttype, rows, context, ti in tables_data:
        if ttype == 'personnel_detail':
            kv = extract_kv_table(rows)

            # Determine role from context — use last matching line (closest to table)
            role = None
            role_keywords = [
                '项目经理', '项目负责人', '系统架构师', '架构师',
                '需求分析师', '需求分析', '测试', '实施', '运维',
                '开发', '前端', '后端', '移动端', '培训',
            ]
            if context:
                for line in reversed(context.split('\n')):
                    for rk in role_keywords:
                        if rk in line:
                            role = rk
                            break
                    if role:
                        break

            name_val = kv.get('姓名', '')
            person = {
                'name': clean_value(name_val),
                'role': role,
                'education': clean_value(kv.get('学历', '')),
                'certifications': [],
                'work_years': clean_value(kv.get('工作年限', '')),
                'industry_experience': None,
                'responsibilities': clean_value(kv.get('角色职责', kv.get('职责', ''))),
                'commitment_ratio': clean_value(kv.get('投入比例', '')),
                'source_table_index': ti,
            }

            # Extract certifications
            for k, v in kv.items():
                if '职称' in k or '认证' in k or '证书' in k:
                    cv = clean_value(v)
                    if cv:
                        # Split on common separators
                        certs = re.split(r'[、，,；;]', cv)
                        person['certifications'] = [c.strip() for c in certs if c.strip()]

            # Industry experience
            for k, v in kv.items():
                if '行业经验' in k or '医疗' in k and '经验' in k:
                    person['industry_experience'] = clean_value(v)

            if is_placeholder(name_val):
                placeholders.append({
                    'field': f'personnel[{len(personnel)}].name',
                    'original': name_val, 'table_index': ti
                })

            personnel.append(person)

        elif ttype == 'team_summary':
            # Multi-column team summary — use header to find correct columns
            header = rows[0]
            col_map = {}
            for ci, h in enumerate(header):
                h = h.strip()
                if '岗位' in h or '角色' in h:
                    col_map['position'] = ci
                elif '人数' in h:
                    col_map['count'] = ci
                elif '职责' in h:
                    col_map['responsibilities'] = ci
                elif '要求' in h:
                    col_map['requirements'] = ci
            for row in rows[1:]:
                entry = {
                    'position': clean_value(row[col_map['position']]) if 'position' in col_map and col_map['position'] < len(row) else None,
                    'count': clean_value(row[col_map['count']]) if 'count' in col_map and col_map['count'] < len(row) else None,
                    'responsibilities': clean_value(row[col_map['responsibilities']]) if 'responsibilities' in col_map and col_map['responsibilities'] < len(row) else None,
                    'requirements': clean_value(row[col_map['requirements']]) if 'requirements' in col_map and col_map['requirements'] < len(row) else None,
                }
                if any(v for v in entry.values()):
                    team_summary.append(entry)

        elif ttype == 'team_summary_single':
            # 2-column single team entry (人数/职责/要求)
            kv = extract_kv_table(rows)
            role = None
            role_keywords = [
                '项目经理', '架构师', '需求分析', '测试', '实施',
                '运维', '开发', '前端', '后端', '培训',
            ]
            if context:
                # Use the LAST paragraph line that contains a role keyword
                # (context lines closer to the table are more relevant)
                for line in reversed(context.split('\n')):
                    for rk in role_keywords:
                        if rk in line:
                            role = rk
                            break
                    if role:
                        break

            entry = {
                'position': role,
                'count': clean_value(kv.get('人数', '')),
                'responsibilities': clean_value(kv.get('角色职责', kv.get('职责', ''))),
                'requirements': clean_value(kv.get('要求', '')),
            }
            team_summary.append(entry)

    return {'personnel': personnel, 'team_summary': team_summary}, placeholders


def extract_performance(tables_data: list[tuple]) -> dict:
    """Extract performance/project history from classified tables."""
    projects = []
    placeholders = []

    for ttype, rows, context, ti in tables_data:
        if ttype != 'performance':
            continue

        header = rows[0]
        # Build column mapping
        col_map = {}
        for ci, h in enumerate(header):
            if '项目名称' in h:
                col_map['project_name'] = ci
            elif '采购人' in h or '甲方' in h:
                col_map['client'] = ci
            elif '合同金额' in h or '金额' in h:
                col_map['contract_amount'] = ci
            elif '签订时间' in h or '合同时间' in h:
                col_map['contract_date'] = ci
            elif '概述' in h or '内容' in h or '说明' in h:
                col_map['description'] = ci
            elif '佐证' in h or '材料' in h:
                col_map['evidence_type'] = ci

        for ri, row in enumerate(rows[1:], 1):
            proj = {}
            all_placeholder = True
            for field, ci in col_map.items():
                if ci < len(row):
                    val = row[ci]
                    if is_placeholder(val):
                        proj[field] = None
                        placeholders.append({
                            'field': f'projects[{len(projects)}].{field}',
                            'original': val, 'table_index': ti
                        })
                    else:
                        cv = clean_value(val)
                        proj[field] = cv
                        if cv:
                            all_placeholder = False

            # Only add if at least one real value exists (or keep structure anyway)
            proj.setdefault('project_name', None)
            proj.setdefault('client', None)
            proj.setdefault('contract_amount', None)
            proj.setdefault('contract_date', None)
            proj.setdefault('description', None)
            proj['image_ref'] = None
            projects.append(proj)

    return {'projects': projects}, placeholders


def extract_qualifications(tables_data: list[tuple], paragraphs: list[str]) -> dict:
    """Extract qualification/certification info."""
    certifications = []
    software_copyrights = []
    placeholders = []

    for ttype, rows, context, ti in tables_data:
        if ttype == 'qualification':
            header = rows[0]
            for row in rows[1:]:
                cert = {
                    'name': clean_value(row[0]) if row else None,
                    'cert_number': clean_value(row[1]) if len(row) > 1 else None,
                    'valid_from': None, 'valid_until': None,
                    'issuing_authority': None, 'scope': None,
                    'image_ref': None,
                }
                certifications.append(cert)

    # Also scan paragraphs for certification mentions
    cert_keywords = {
        'ISO 9001': '质量管理体系', 'ISO9001': '质量管理体系',
        'ISO 14001': '环境管理体系', 'ISO14001': '环境管理体系',
        'ISO 27001': '信息安全管理体系', 'ISO27001': '信息安全管理体系',
        'ISO 20000': '信息技术服务管理体系', 'ISO20000': '信息技术服务管理体系',
        'ISO 45001': '职业健康安全管理体系', 'OHSAS18001': '职业健康安全管理体系',
        'CMMI': '能力成熟度模型集成',
        '软件著作权': '软件著作权',
        '高新技术企业': '高新技术企业认定',
    }
    found_certs = set()
    for para in paragraphs:
        for kw, desc in cert_keywords.items():
            if kw in para and kw not in found_certs:
                found_certs.add(kw)
                # Check if already in certifications list
                already = any(c['name'] and kw.replace(' ', '') in c['name'].replace(' ', '')
                             for c in certifications)
                if not already:
                    certifications.append({
                        'name': f'{kw} ({desc})',
                        'cert_number': None, 'valid_from': None,
                        'valid_until': None, 'issuing_authority': None,
                        'scope': None, 'image_ref': None,
                        '_source': 'paragraph_scan',
                    })

    return {'certifications': certifications, 'software_copyrights': software_copyrights}, placeholders


def extract_cost_breakdown(tables_data: list[tuple]) -> dict:
    """Extract cost breakdown information."""
    result = {
        'construction_costs': [],
        'warranty_annual_costs': [],
        'post_warranty_annual_costs': [],
        'summary': [],
    }
    placeholders = []

    cost_tables = [(rows, context, ti) for ttype, rows, context, ti in tables_data
                   if ttype == 'cost_breakdown']

    for idx, (rows, context, ti) in enumerate(cost_tables):
        header = rows[0]
        ncols = len(header)

        # Determine which cost category based on context and header
        category = 'construction_costs'
        ctx = context.lower() if context else ''
        if '质保' in ctx and '后' not in ctx:
            category = 'warranty_annual_costs'
        elif '质保' in ctx and '后' in ctx:
            category = 'post_warranty_annual_costs'
        elif '汇总' in ctx or '合计' in header[0]:
            category = 'summary'
        elif '年费用' in ' '.join(header):
            # Has annual cost columns - likely warranty or post-warranty
            if '质保' in ctx or idx == 1:
                category = 'warranty_annual_costs'
            else:
                category = 'post_warranty_annual_costs'

        for row in rows[1:]:
            if not any(clean_value(c) for c in row):
                continue
            if ncols >= 4 and '年费用' in ' '.join(header):
                entry = {
                    'item': clean_value(row[0]) if row else None,
                    'annual': clean_value(row[1]) if len(row) > 1 else None,
                    'total': clean_value(row[2]) if len(row) > 2 else None,
                    'note': clean_value(row[3]) if len(row) > 3 else None,
                }
            elif '阶段' in header[0] or '占比' in ' '.join(header):
                entry = {
                    'phase': clean_value(row[0]) if row else None,
                    'amount': clean_value(row[1]) if len(row) > 1 else None,
                    'percentage': clean_value(row[2]) if len(row) > 2 else None,
                    'note': clean_value(row[3]) if len(row) > 3 else None,
                }
                category = 'summary'
            else:
                entry = {
                    'item': clean_value(row[0]) if row else None,
                    'amount': clean_value(row[1]) if len(row) > 1 else None,
                    'note': clean_value(row[2]) if len(row) > 2 else None,
                }

            result[category].append(entry)

    return result, placeholders


def cross_reference_index(data: dict, index_path: str) -> dict:
    """Cross-reference extracted text data with image index."""
    if not os.path.exists(index_path):
        return data

    with open(index_path, 'r', encoding='utf-8') as f:
        index = json.load(f)

    documents = index.get('documents', [])

    # Cross-reference qualifications with certification images
    if 'qualifications' in data:
        for cert in data['qualifications'].get('certifications', []):
            if not cert.get('name'):
                continue
            name = cert['name'].lower()
            for doc_entry in documents:
                if doc_entry['category'] != '资质证明':
                    continue
                tags = ' '.join(doc_entry.get('searchable_tags', [])).lower()
                doc_type = doc_entry.get('type', '').lower()
                # Simple keyword matching
                if any(kw in tags or kw in doc_type for kw in name.split()):
                    cert['image_ref'] = doc_entry['id']
                    break

    # Cross-reference performance with contract images
    if 'performance' in data:
        contract_docs = [d for d in documents if d['category'] == '业绩证明']
        for pi, proj in enumerate(data['performance'].get('projects', [])):
            if pi < len(contract_docs):
                proj['image_ref'] = contract_docs[pi]['id']

    return data


def extract_all(docx_path: str, output_dir: str, index_path: str = None):
    """Main extraction function."""
    os.makedirs(output_dir, exist_ok=True)

    print(f"Loading {docx_path}...")
    doc = Document(docx_path)
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # Get paragraph contexts for each table
    contexts = get_table_contexts(doc)

    # Classify all tables
    classified = []
    type_counts = {}
    for ti, table in enumerate(doc.tables):
        rows = table_to_rows(table)
        context = contexts.get(ti, '')
        ttype = classify_table(rows, context)
        classified.append((ttype, rows, context, ti))
        type_counts[ttype] = type_counts.get(ttype, 0) + 1

    print(f"\nTable classification:")
    for t, c in sorted(type_counts.items(), key=lambda x: -x[1]):
        print(f"  {t}: {c}")

    # Extract paragraphs text for qualification scanning
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract each data type
    all_placeholders = []
    source_file = os.path.basename(docx_path)
    timestamp = datetime.now().isoformat()

    # Company
    print("\nExtracting company info...")
    company, ph = extract_company(classified)
    company['source_file'] = source_file
    company['extracted_at'] = timestamp
    all_placeholders.extend(ph)
    _save_json(os.path.join(output_dir, 'company.json'), company)
    _print_company_summary(company)

    # Personnel
    print("\nExtracting personnel info...")
    personnel, ph = extract_personnel(classified)
    personnel['source_file'] = source_file
    personnel['extracted_at'] = timestamp
    all_placeholders.extend(ph)
    _save_json(os.path.join(output_dir, 'personnel.json'), personnel)
    print(f"  {len(personnel['personnel'])} named personnel, {len(personnel['team_summary'])} team entries")

    # Performance
    print("\nExtracting performance/projects...")
    performance, ph = extract_performance(classified)
    performance['source_file'] = source_file
    performance['extracted_at'] = timestamp
    all_placeholders.extend(ph)
    _save_json(os.path.join(output_dir, 'performance.json'), performance)
    print(f"  {len(performance['projects'])} projects")

    # Qualifications
    print("\nExtracting qualifications...")
    qualifications, ph = extract_qualifications(classified, paragraphs)
    qualifications['source_file'] = source_file
    qualifications['extracted_at'] = timestamp
    all_placeholders.extend(ph)
    _save_json(os.path.join(output_dir, 'qualifications.json'), qualifications)
    print(f"  {len(qualifications['certifications'])} certifications")

    # Cost breakdown
    print("\nExtracting cost breakdown...")
    cost, ph = extract_cost_breakdown(classified)
    cost['source_file'] = source_file
    cost['extracted_at'] = timestamp
    all_placeholders.extend(ph)
    _save_json(os.path.join(output_dir, 'cost_breakdown.json'), cost)

    # Cross-reference with image index if available
    if index_path and os.path.exists(index_path):
        print(f"\nCross-referencing with {index_path}...")
        all_data = {
            'qualifications': qualifications,
            'performance': performance,
        }
        cross_reference_index(all_data, index_path)
        # Re-save with image refs
        _save_json(os.path.join(output_dir, 'qualifications.json'), qualifications)
        _save_json(os.path.join(output_dir, 'performance.json'), performance)

    # Save placeholder summary
    if all_placeholders:
        _save_json(os.path.join(output_dir, '_placeholders.json'), {
            'count': len(all_placeholders),
            'items': all_placeholders,
            'source_file': source_file,
        })
        print(f"\n{len(all_placeholders)} placeholders found (saved to _placeholders.json)")

    # Generate extraction report
    _generate_report(output_dir, company, personnel, performance,
                     qualifications, cost, all_placeholders, source_file)

    print(f"\nDone. Output: {output_dir}/")


def _save_json(path: str, data: dict):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _print_company_summary(company: dict):
    name = company.get('name') or '(未提取)'
    credit = company.get('credit_code') or '(未提取)'
    legal = company.get('legal_representative', {}).get('name') or '(未提取)'
    auth = company.get('authorized_representative', {}).get('name') or '(未提取)'
    print(f"  公司: {name}")
    print(f"  信用代码: {credit}")
    print(f"  法人: {legal}")
    print(f"  授权代表: {auth}")


def _generate_report(output_dir, company, personnel, performance,
                     qualifications, cost, placeholders, source_file):
    """Generate a markdown extraction report."""
    lines = [
        '# 资料提取报告',
        '',
        f'- 来源文件: `{source_file}`',
        f'- 提取时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}',
        '',
        '## 提取统计',
        '',
        '| 类别 | 数量 | 文件 |',
        '|------|------|------|',
        f'| 公司信息 | 1 | company.json |',
        f'| 人员详情 | {len(personnel["personnel"])} | personnel.json |',
        f'| 团队汇总 | {len(personnel["team_summary"])} | personnel.json |',
        f'| 业绩项目 | {len(performance["projects"])} | performance.json |',
        f'| 资质证书 | {len(qualifications["certifications"])} | qualifications.json |',
        f'| 成本明细 | {sum(len(v) for v in cost.values() if isinstance(v, list))} | cost_breakdown.json |',
        '',
        '## 公司信息摘要',
        '',
        f'- 名称: {company.get("name") or "❌ 未提取"}',
        f'- 信用代码: {company.get("credit_code") or "❌ 未提取"}',
        f'- 注册地址: {company.get("registered_address") or "❌ 未提取"}',
        f'- 法人: {company.get("legal_representative", {}).get("name") or "❌ 未提取"}',
        f'- 授权代表: {company.get("authorized_representative", {}).get("name") or "❌ 未提取"}',
        '',
    ]

    if personnel['personnel']:
        lines.extend([
            '## 人员信息',
            '',
            '| 姓名 | 角色 | 学历 | 资质 |',
            '|------|------|------|------|',
        ])
        for p in personnel['personnel']:
            name = p.get('name') or '(待补充)'
            role = p.get('role') or '—'
            edu = p.get('education') or '—'
            certs = '、'.join(p.get('certifications', [])) or '—'
            lines.append(f'| {name} | {role} | {edu} | {certs} |')
        lines.append('')

    if placeholders:
        lines.extend([
            '## 未填充占位符',
            '',
            f'共 {len(placeholders)} 个占位符待补充：',
            '',
            '| 字段 | 原始值 | 表格序号 |',
            '|------|--------|---------|',
        ])
        for ph in placeholders:
            lines.append(f'| {ph["field"]} | {ph["original"]} | {ph["table_index"]} |')
        lines.append('')

    report_path = os.path.join(output_dir, 'extraction_report.md')
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"  Report: {report_path}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Extract structured text from bid response Word file')
    parser.add_argument('docx_path', help='Path to the .docx file')
    parser.add_argument('--output-dir', default='data', help='Output directory for JSON files')
    parser.add_argument('--index', default=None, help='Path to existing index.json for cross-referencing')
    args = parser.parse_args()
    extract_all(args.docx_path, args.output_dir, args.index)
